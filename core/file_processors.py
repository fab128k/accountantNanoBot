# core/file_processors.py
# AccountantNanoBot v1.0.0 - Processori per file uploadati
# ============================================================================
#
# Tipi supportati:
# - Documenti: PDF, TXT, MD, DOCX, XML (incluso FatturaPA)
# - Immagini: PNG, JPG, JPEG (solo modelli Vision)
#
# ============================================================================

import io
import base64
from pathlib import Path
from typing import Optional, Tuple, List
from dataclasses import dataclass, field

from config import MAX_DOCUMENT_CHARS


@dataclass
class ProcessedFile:
    """
    Risultato del processamento di un file uploadato.

    Attributes:
        filename: Nome originale del file
        file_type: "document" | "image" | "xml_fattura" | "unknown"
        content: Testo estratto (documenti) o base64 (immagini)
        mime_type: MIME type del file
        size_bytes: Dimensione in bytes
        preview: Anteprima per UI
        error: Messaggio di errore (se presente)
    """
    filename: str
    file_type: str
    content: str
    mime_type: str
    size_bytes: int
    preview: str = ""
    error: Optional[str] = None


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def is_image_file(filename: str) -> bool:
    ext = Path(filename).suffix.lower()
    return ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"]


def is_document_file(filename: str) -> bool:
    ext = Path(filename).suffix.lower()
    return ext in [".pdf", ".txt", ".md", ".docx"]


def is_xml_file(filename: str) -> bool:
    ext = Path(filename).suffix.lower()
    return ext == ".xml"


def get_file_extension(filename: str) -> str:
    return Path(filename).suffix.lower()


# ============================================================================
# DOCUMENT EXTRACTORS
# ============================================================================

def extract_text_from_txt(file_bytes: bytes, encoding: str = "utf-8") -> str:
    try:
        return file_bytes.decode(encoding)
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception:
            return file_bytes.decode("utf-8", errors="ignore")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []

        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(f"[Pagina {page_num}]\n{page_text.strip()}")

        if not text_parts:
            return "[PDF senza testo estraibile - potrebbe contenere solo immagini]"

        return "\n\n".join(text_parts)

    except ImportError:
        return "[Errore: PyPDF2 non installato]"
    except Exception as e:
        return f"[Errore estrazione PDF: {e}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    table_rows.append(" | ".join(row_cells))
            if table_rows:
                paragraphs.append("\n".join(table_rows))

        if not paragraphs:
            return "[Documento DOCX vuoto]"

        return "\n\n".join(paragraphs)

    except ImportError:
        return "[Errore: python-docx non installato]"
    except Exception as e:
        return f"[Errore estrazione DOCX: {e}]"


def extract_text_from_xml(file_bytes: bytes, filename: str = "") -> str:
    """
    Estrae testo da file XML.

    Se il root element è FatturaElettronica, delega al parser FatturaPA
    per ottenere un testo strutturato leggibile.
    Altrimenti esegue un pretty-print del contenuto XML.

    Args:
        file_bytes: Contenuto del file XML in bytes
        filename: Nome del file (per log)

    Returns:
        Testo estratto/formattato
    """
    try:
        from lxml import etree

        root = etree.fromstring(file_bytes)

        # Controlla se è una FatturaPA
        local_name = etree.QName(root.tag).localname if root.tag.startswith('{') else root.tag

        if local_name == "FatturaElettronica":
            # Delega al parser specializzato
            try:
                from parsers.fattura_pa import FatturaPAParser
                parser = FatturaPAParser()
                fatture = parser.parse_bytes(file_bytes)
                if fatture:
                    parts = [parser.to_text_summary(f) for f in fatture]
                    return "\n\n---\n\n".join(parts)
            except Exception as e:
                pass  # Fallback a XML generico

        # Pretty-print XML generico
        pretty = etree.tostring(root, pretty_print=True, encoding="unicode")
        return pretty[:MAX_DOCUMENT_CHARS] if len(pretty) > MAX_DOCUMENT_CHARS else pretty

    except ImportError:
        # Fallback senza lxml
        try:
            import xml.etree.ElementTree as ET
            root = ET.fromstring(file_bytes)
            return ET.tostring(root, encoding="unicode")
        except Exception as e:
            return f"[Errore parsing XML: {e}]"
    except Exception as e:
        return f"[Errore estrazione XML: {e}]"


# ============================================================================
# IMAGE PROCESSING
# ============================================================================

def process_image_to_base64(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    ext = get_file_extension(filename)

    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }

    mime_type = mime_map.get(ext, "image/png")
    base64_data = base64.b64encode(file_bytes).decode("utf-8")

    return base64_data, mime_type


def create_image_thumbnail(file_bytes: bytes, max_size: int = 200) -> Optional[str]:
    try:
        from PIL import Image

        img = Image.open(io.BytesIO(file_bytes))

        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = background

        img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)

        buffer = io.BytesIO()
        img.save(buffer, format="PNG", optimize=True)
        buffer.seek(0)

        return base64.b64encode(buffer.getvalue()).decode("utf-8")

    except (ImportError, Exception):
        return None


# ============================================================================
# MAIN PROCESSOR
# ============================================================================

def process_uploaded_file(uploaded_file) -> ProcessedFile:
    """
    Processa un file uploadato da Streamlit.
    """
    filename = uploaded_file.name
    file_bytes = uploaded_file.read()
    size_bytes = len(file_bytes)
    ext = get_file_extension(filename)

    uploaded_file.seek(0)

    # ========== IMMAGINI ==========
    if is_image_file(filename):
        base64_data, mime_type = process_image_to_base64(file_bytes, filename)
        thumbnail = create_image_thumbnail(file_bytes)

        return ProcessedFile(
            filename=filename,
            file_type="image",
            content=base64_data,
            mime_type=mime_type,
            size_bytes=size_bytes,
            preview=thumbnail if thumbnail else "",
            error=None
        )

    # ========== XML (incluso FatturaPA) ==========
    if is_xml_file(filename):
        text = extract_text_from_xml(file_bytes, filename)
        error = None
        if text.startswith("[Errore"):
            error = text

        original_len = len(text)
        if original_len > MAX_DOCUMENT_CHARS:
            text = text[:MAX_DOCUMENT_CHARS] + f"\n\n[... troncato a {MAX_DOCUMENT_CHARS:,} caratteri]"

        preview = text[:500] + "..." if len(text) > 500 else text

        return ProcessedFile(
            filename=filename,
            file_type="xml_fattura" if "FatturaElettronica" in text[:200] else "document",
            content=text,
            mime_type="application/xml",
            size_bytes=size_bytes,
            preview=preview,
            error=error
        )

    # ========== DOCUMENTI ==========
    if is_document_file(filename):
        if ext == ".pdf":
            text = extract_text_from_pdf(file_bytes)
            mime_type = "application/pdf"
        elif ext == ".docx":
            text = extract_text_from_docx(file_bytes)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif ext == ".md":
            text = extract_text_from_txt(file_bytes)
            mime_type = "text/markdown"
        else:
            text = extract_text_from_txt(file_bytes)
            mime_type = "text/plain"

        error = None
        if text.startswith("[Errore"):
            error = text

        original_len = len(text)
        if original_len > MAX_DOCUMENT_CHARS:
            text = text[:MAX_DOCUMENT_CHARS] + f"\n\n[... troncato a {MAX_DOCUMENT_CHARS:,} caratteri su {original_len:,} totali]"

        preview = text[:500] + "..." if len(text) > 500 else text

        return ProcessedFile(
            filename=filename,
            file_type="document",
            content=text,
            mime_type=mime_type,
            size_bytes=size_bytes,
            preview=preview,
            error=error
        )

    # ========== TIPO NON SUPPORTATO ==========
    return ProcessedFile(
        filename=filename,
        file_type="unknown",
        content="",
        mime_type="application/octet-stream",
        size_bytes=size_bytes,
        preview="",
        error=f"Tipo file non supportato: {ext}"
    )


def process_multiple_files(uploaded_files: List) -> List[ProcessedFile]:
    processed = []
    for f in uploaded_files:
        if f is not None:
            processed.append(process_uploaded_file(f))
    return processed


def build_document_context(processed_files: List[ProcessedFile]) -> str:
    documents = [f for f in processed_files if f.file_type in ("document", "xml_fattura") and not f.error]

    if not documents:
        return ""

    parts = []
    for doc in documents:
        parts.append(f"[📄 File: {doc.filename}]\n{doc.content}")

    return "\n\n--- FILE ALLEGATO ---\n".join(parts)


def get_images_for_vision(processed_files: List[ProcessedFile]) -> List[dict]:
    images = [f for f in processed_files if f.file_type == "image" and not f.error]

    return [
        {
            "filename": img.filename,
            "base64": img.content,
            "mime_type": img.mime_type
        }
        for img in images
    ]


def get_attachment_names(processed_files: List[ProcessedFile]) -> List[str]:
    return [f.filename for f in processed_files if not f.error]
